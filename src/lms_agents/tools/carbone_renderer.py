"""
Carbone.io PDF Renderer — generates professional worksheet PDFs.

Workflow:
1. Build a DOCX template with Carbone placeholders using python-docx
2. Upload template to Carbone API → template_id (cached per theme)
3. POST content data to render → render_id
4. GET rendered PDF bytes

Requires CARBONE_API_KEY env var. Falls back gracefully if unavailable.
"""
import io
import logging
import os
from typing import Optional

import httpx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

log = logging.getLogger(__name__)

CARBONE_BASE = "https://api.carbone.io"
CARBONE_TIMEOUT = 30.0

THEME_COLORS = {
    "modern_clean": {"primary": "F97316", "accent": "FB923C", "bg": "FFF7ED"},
    "ocean_blue": {"primary": "2563EB", "accent": "3B82F6", "bg": "EFF6FF"},
    "forest_green": {"primary": "059669", "accent": "10B981", "bg": "ECFDF5"},
    "royal_purple": {"primary": "7C3AED", "accent": "8B5CF6", "bg": "F5F3FF"},
}

# Module-level cache: theme_name → template_id
_template_cache: dict[str, str] = {}


def _hex_to_rgb(hex_str: str) -> RGBColor:
    """Convert '2563EB' → RGBColor(0x25, 0x63, 0xEB)."""
    return RGBColor(int(hex_str[:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def _carbone_headers() -> dict:
    api_key = os.environ.get("CARBONE_API_KEY", "")
    return {
        "Authorization": f"Bearer {api_key}",
        "carbone-version": "4",
    }


# ── Carbone API functions ─────────────────────────────────────────────────


def upload_template(template_bytes: bytes, template_name: str) -> str:
    """Upload a DOCX template to Carbone. Returns template_id."""
    url = f"{CARBONE_BASE}/template"
    files = {"template": (template_name, template_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    resp = httpx.post(url, headers=_carbone_headers(), files=files, timeout=CARBONE_TIMEOUT)
    resp.raise_for_status()
    body = resp.json()
    if not body.get("success"):
        raise RuntimeError(f"Carbone upload failed: {body}")
    template_id = body["data"]["templateId"]
    log.info(f"[Carbone] Uploaded template '{template_name}' → {template_id}")
    return template_id


def render_document(template_id: str, data: dict, convert_to: str = "pdf") -> bytes:
    """Render a template with data and return the output file bytes."""
    headers = {**_carbone_headers(), "Content-Type": "application/json"}

    # Step 1: request render
    render_url = f"{CARBONE_BASE}/render/{template_id}"
    render_resp = httpx.post(
        render_url, headers=headers,
        json={"data": data, "convertTo": convert_to},
        timeout=CARBONE_TIMEOUT,
    )
    render_resp.raise_for_status()
    render_body = render_resp.json()
    if not render_body.get("success"):
        raise RuntimeError(f"Carbone render failed: {render_body}")
    render_id = render_body["data"]["renderId"]

    # Step 2: download rendered file
    dl_url = f"{CARBONE_BASE}/render/{render_id}"
    dl_resp = httpx.get(dl_url, headers=_carbone_headers(), timeout=CARBONE_TIMEOUT)
    dl_resp.raise_for_status()
    log.info(f"[Carbone] Rendered {len(dl_resp.content)} bytes (template={template_id})")
    return dl_resp.content


# ── Template builder (python-docx) ────────────────────────────────────────


def _set_cell_shading(cell, hex_color: str):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {qn("w:fill"): hex_color, qn("w:val"): "clear"},
    )
    shading.append(shading_elem)


def _add_border_to_paragraph(paragraph, color_hex: str):
    """Add a bottom border to a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    borders = pPr.makeelement(qn("w:pBdr"), {})
    bottom = borders.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): color_hex,
        },
    )
    borders.append(bottom)
    pPr.append(borders)


def _create_worksheet_template(theme: str) -> bytes:
    """
    Build a professional DOCX template with Carbone placeholders.
    Returns the DOCX as bytes.
    """
    colors = THEME_COLORS.get(theme, THEME_COLORS["modern_clean"])
    primary = colors["primary"]
    accent = colors["accent"]
    bg = colors["bg"]

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ── Default font ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1C, 0x19, 0x17)
    style.paragraph_format.space_after = Pt(4)

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(2)
    title_run = title_para.add_run("{d.title}")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = _hex_to_rgb(primary)
    title_run.font.name = "Calibri"

    # ── Accent line under title ──
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_border_to_paragraph(line_para, accent)
    line_para.paragraph_format.space_after = Pt(8)

    # ── Name / Date line ──
    nd_para = doc.add_paragraph()
    nd_para.paragraph_format.space_before = Pt(4)
    nd_para.paragraph_format.space_after = Pt(12)
    nd_run = nd_para.add_run("Name: ________________________________     Date: ________________")
    nd_run.font.size = Pt(11)
    nd_run.font.color.rgb = RGBColor(0x57, 0x53, 0x4E)

    # ── Instructions ──
    instr_para = doc.add_paragraph()
    instr_para.paragraph_format.space_after = Pt(10)
    instr_run = instr_para.add_run("{d.instructions}")
    instr_run.italic = True
    instr_run.font.size = Pt(10.5)
    instr_run.font.color.rgb = RGBColor(0x44, 0x40, 0x3C)

    # ── Word Bank (table with one row, items inline) ──
    wb_header = doc.add_paragraph()
    wb_header.paragraph_format.space_before = Pt(8)
    wb_header.paragraph_format.space_after = Pt(4)
    wb_run = wb_header.add_run("Word Bank")
    wb_run.bold = True
    wb_run.font.size = Pt(12)
    wb_run.font.color.rgb = _hex_to_rgb(primary)

    # Carbone iterates over word_bank array — one cell per word in a table
    wb_table = doc.add_table(rows=1, cols=1)
    wb_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = wb_table.cell(0, 0)
    _set_cell_shading(cell, bg)
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_para.add_run("{d.word_bank[i].value}")
    run.font.size = Pt(11)
    run.font.color.rgb = _hex_to_rgb(primary)
    run.bold = True

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

    # ── Questions section header ──
    q_header = doc.add_paragraph()
    q_header.paragraph_format.space_before = Pt(10)
    q_header.paragraph_format.space_after = Pt(6)
    _add_border_to_paragraph(q_header, accent)
    q_run = q_header.add_run("Questions")
    q_run.bold = True
    q_run.font.size = Pt(14)
    q_run.font.color.rgb = _hex_to_rgb(primary)

    # ── Question template (Carbone iterates d.questions) ──
    q_para = doc.add_paragraph()
    q_para.paragraph_format.space_after = Pt(2)
    q_num = q_para.add_run("{d.questions[i].number}. ")
    q_num.bold = True
    q_num.font.size = Pt(11)
    q_text = q_para.add_run("{d.questions[i].question_text}")
    q_text.font.size = Pt(11)

    # Multiple choice options (Carbone conditional — always present, hidden if empty)
    for idx, letter in enumerate(["A", "B", "C", "D"]):
        opt_para = doc.add_paragraph()
        opt_para.paragraph_format.left_indent = Cm(1)
        opt_para.paragraph_format.space_after = Pt(1)
        opt_run = opt_para.add_run(f"     {letter})  {{d.questions[i].options[{idx}]}}")
        opt_run.font.size = Pt(10.5)
        opt_run.font.color.rgb = RGBColor(0x57, 0x53, 0x4E)

    # Answer line (for short_answer / fill_in_blank)
    ans_para = doc.add_paragraph()
    ans_para.paragraph_format.space_before = Pt(4)
    ans_para.paragraph_format.space_after = Pt(12)
    ans_run = ans_para.add_run("     Answer: _______________________________________________")
    ans_run.font.size = Pt(10.5)
    ans_run.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)

    # ── Answer Key section ──
    ak_header = doc.add_paragraph()
    ak_header.paragraph_format.space_before = Pt(16)
    ak_header.paragraph_format.space_after = Pt(6)
    _add_border_to_paragraph(ak_header, accent)
    ak_run = ak_header.add_run("Answer Key")
    ak_run.bold = True
    ak_run.font.size = Pt(14)
    ak_run.font.color.rgb = _hex_to_rgb(primary)

    ak_para = doc.add_paragraph()
    ak_para.paragraph_format.space_after = Pt(2)
    ak_num = ak_para.add_run("{d.answer_key[i].number}. ")
    ak_num.bold = True
    ak_num.font.size = Pt(10)
    ak_ans = ak_para.add_run("{d.answer_key[i].answer}")
    ak_ans.font.size = Pt(10)
    ak_ans.font.color.rgb = _hex_to_rgb(accent)

    # ── Serialize to bytes ──
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── High-level render function ─────────────────────────────────────────────


def _prepare_carbone_data(content: dict) -> dict:
    """
    Transform the flat worksheet content dict into Carbone-compatible format.

    - word_bank: list[str] → list[{value: str}]
    - questions: ensure options is always a 4-element list
    - answer_key: ensure each has number + answer
    """
    # Word bank: convert flat strings to objects
    raw_wb = content.get("word_bank", [])
    if raw_wb and isinstance(raw_wb[0], str):
        word_bank = [{"value": w} for w in raw_wb]
    else:
        word_bank = raw_wb  # already objects

    # Questions: normalize options to always have 4 elements
    questions = []
    for q in content.get("questions", []):
        opts = q.get("options") or []
        # Pad to exactly 4 options (Carbone template expects 4 slots)
        while len(opts) < 4:
            opts.append("")
        entry = {
            "number": q.get("number", 0),
            "question_text": q.get("question_text", ""),
            "type": q.get("type", "short_answer"),
            "options": opts[:4],
            "correct_answer": q.get("correct_answer", ""),
            "points": q.get("points", 1),
        }
        questions.append(entry)

    # Answer key
    answer_key = []
    for ak in content.get("answer_key", []):
        answer_key.append({
            "number": ak.get("number", 0),
            "answer": str(ak.get("answer", "")),
        })

    return {
        "title": content.get("title", "Worksheet"),
        "instructions": content.get("instructions", ""),
        "word_bank": word_bank,
        "questions": questions,
        "answer_key": answer_key,
    }


def render_worksheet(content: dict, theme: str = "modern_clean") -> bytes:
    """
    End-to-end: build template → upload to Carbone → render PDF → return bytes.

    Uses cached template_id per theme to avoid re-uploading each time.
    Raises on failure (caller should handle fallback).
    """
    api_key = os.environ.get("CARBONE_API_KEY")
    if not api_key:
        raise RuntimeError("CARBONE_API_KEY not configured")

    # Ensure valid theme
    if theme not in THEME_COLORS:
        theme = "modern_clean"

    # Get or create cached template_id
    template_id = _template_cache.get(theme)
    if not template_id:
        log.info(f"[Carbone] Building DOCX template for theme '{theme}'")
        docx_bytes = _create_worksheet_template(theme)
        template_id = upload_template(docx_bytes, f"worksheet_{theme}.docx")
        _template_cache[theme] = template_id

    # Prepare data for Carbone
    data = _prepare_carbone_data(content)

    # Render PDF
    pdf_bytes = render_document(template_id, data, convert_to="pdf")
    log.info(f"[Carbone] Worksheet PDF rendered: {len(pdf_bytes)} bytes, theme={theme}")
    return pdf_bytes
